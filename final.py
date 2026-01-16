import streamlit as st
import torch
from sentence_transformers import SentenceTransformer, util
import pandas as pd
import speech_recognition as sr
import os
import re
import pytesseract
from docx import Document
from PIL import Image
import string
from googletrans import Translator

pytesseract.pytesseract.tesseract_cmd = r"D:\legal_project\tesseract\tesseract.exe"

MODEL_PATH = "legal_nlp_model"
model = SentenceTransformer(MODEL_PATH)
knowledge_base = pd.read_csv("Final_QA_dataset.csv")
corpus_embeddings = torch.load("corpus_embeddings.pt")
translator = Translator()

def normalize(text):
    return text.lower().translate(str.maketrans('', '', string.punctuation)).strip()

def translate_text(text, src='auto', dest='en'):
    return translator.translate(text, src=src, dest=dest).text

def get_best_legal_answer(query):
    query_clean = normalize(query)
    for _, row in knowledge_base.iterrows():
        if normalize(row['Question']) == query_clean:
            return row['Answer'], row['category'], "Exact Match"
    query_embedding = model.encode(query, convert_to_tensor=True)
    scores = util.pytorch_cos_sim(query_embedding, corpus_embeddings)[0]
    best_idx = torch.argmax(scores).item()
    return knowledge_base.iloc[best_idx]['Answer'], knowledge_base.iloc[best_idx]['category']

def get_word_template(template_name, folder_path="word_templates"):
    for filename in os.listdir(folder_path):
        if template_name.lower() in filename.lower() and filename.endswith(".docx"):
            return os.path.join(folder_path, filename)
    return None

def get_template_images(template_name, folder_path="templates_img"):
    return [os.path.join(folder_path, f) for f in os.listdir(folder_path)
            if template_name.lower() in f.lower() and f.lower().endswith('.jpg')]

def extract_fields(template_path):
    doc = Document(template_path)
    fields = set()
    pattern = r'\{\{([^}]+)\}\}'
    for paragraph in doc.paragraphs:
        fields.update(match.strip() for match in re.findall(pattern, paragraph.text))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                fields.update(match.strip() for match in re.findall(pattern, cell.text))
    return sorted(fields)

def fill_template(template_path, user_inputs):
    doc = Document(template_path)
    for paragraph in doc.paragraphs:
        for field, value in user_inputs.items():
            placeholder = f"{{{{{field}}}}}"
            paragraph.text = paragraph.text.replace(placeholder, value)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for field, value in user_inputs.items():
                    placeholder = f"{{{{{field}}}}}"
                    cell.text = cell.text.replace(placeholder, value)
    return doc

def get_voice_input():
    recognizer = sr.Recognizer()
    with sr.Microphone() as source:
        st.info("🎤 Listening... Speak now.")
        recognizer.adjust_for_ambient_noise(source)
        audio = recognizer.listen(source)
    try:
        return recognizer.recognize_google(audio)
    except:
        return "Sorry, could not process the voice input."

def extract_text_from_file(file):
    file_ext = file.name.split('.')[-1].lower()
    if file_ext == "pdf":
        import pdfplumber
        with pdfplumber.open(file) as pdf:
            return "\n".join([page.extract_text() for page in pdf.pages if page.extract_text()])
    elif file_ext == "docx":
        doc = Document(file)
        return "\n".join([para.text for para in doc.paragraphs])
    elif file_ext in ["jpg", "jpeg", "png"]:
        image = Image.open(file)
        return pytesseract.image_to_string(image)
    elif file_ext == "txt":
        return file.read().decode("utf-8")
    return ""

def main():
    st.title("Legal Digital Assistant")

    tab1, tab2 = st.tabs(["Legal Information", "Document Templates"])

    # --- TAB 1: LEGAL INFO ---
    with tab1:
        st.header("Legal Information")

        languages = {
            "English": "en", "Hindi": "hi", "Tamil": "ta", "Telugu": "te", "Kannada": "kn", "Marathi": "mr", "Bengali": "bn"
        }
        selected_lang_name = st.selectbox("🌐 Select Language", list(languages.keys()), index=0)
        selected_lang = languages[selected_lang_name]

        if "legal_query" not in st.session_state:
            st.session_state.legal_query = ""

        input_tab1, input_tab2 = st.tabs(["🖊️ Text Input", "🎙️ Voice Input"])

        with input_tab1:
            st.session_state.legal_query = st.text_area("Ask your legal question:", value=st.session_state.legal_query)

            st.markdown("<div style='text-align: center;'>— or —</div>", unsafe_allow_html=True)

            uploaded_file = st.file_uploader("Upload a related document or image", type=["pdf", "docx", "txt", "jpg", "jpeg", "png"])
            if uploaded_file:
                extracted_text = extract_text_from_file(uploaded_file)
                st.text_area("Extracted Text:", value=extracted_text, height=200)
                if st.button("Analyze Extracted Text", key="analyze_extracted"):
                    st.session_state.legal_query = extracted_text

        with input_tab2:
            if st.button("Record Voice Query", key="legal_voice"):
                voice_text = get_voice_input()
                st.session_state.legal_query = voice_text
                st.write(voice_text)

        if st.button("Submit", key="submit_query"):
            user_query = st.session_state.legal_query.strip()
            if user_query:
                st.success("Searching legal knowledge base...")

                translated_input = translate_text(user_query, src=selected_lang, dest='en')
                answer, category = get_best_legal_answer(translated_input)

                answer_t = translate_text(answer, src='en', dest=selected_lang)
                category_t = translate_text(category, src='en', dest=selected_lang)

                st.write("**Legal Answer:**", answer_t)
                st.write("**Category:**", category_t)
            else:
                st.warning("Please enter or upload a legal question.")

        # ✅ Reset Query Button
        if st.button("Reset Query", key="reset_query"):
            st.session_state.legal_query = ""

    # --- TAB 2: DOCUMENT TEMPLATES ---
    with tab2:
        st.header("Document Template Generator")
        templates_dir = "word_templates"
        os.makedirs(templates_dir, exist_ok=True)

        template_files = [f.split('.')[0] for f in os.listdir(templates_dir) if f.endswith('.docx')]
        template_selection = st.selectbox("Select a document template:", ["-- Select --"] + template_files)

        if template_selection != "-- Select --":
            template_path = get_word_template(template_selection)
            image_paths = get_template_images(template_selection)

            if image_paths:
                st.subheader("Template Previews:")
                cols = st.columns(len(image_paths))
                for idx, img_path in enumerate(image_paths):
                    with cols[idx]:
                        st.image(img_path, use_column_width=True)
            else:
                st.info("No preview images found for this template.")

            required_fields = extract_fields(template_path)
            if required_fields:
                st.success(f"Found {len(required_fields)} fields to fill.")
            else:
                st.warning("No placeholders found. Make sure to use {{FIELD_NAME}} format.")

            if "edit_mode" not in st.session_state:
                st.session_state.edit_mode = False
            if "user_inputs" not in st.session_state:
                st.session_state.user_inputs = {}

            if st.button("Edit Template", key="edit_button"):
                st.session_state.edit_mode = True
                for field in required_fields:
                    st.session_state.user_inputs.setdefault(field, "")

            if st.session_state.edit_mode and required_fields:
                with st.form("template_form"):
                    for field in required_fields:
                        st.session_state.user_inputs[field] = st.text_input(
                            f"{field}:",
                            value=st.session_state.user_inputs.get(field, ""),
                            key=f"input_{field}"
                        )
                    submitted = st.form_submit_button("Generate Document")
                    if submitted:
                        if all(st.session_state.user_inputs.values()):
                            completed_doc = fill_template(template_path, st.session_state.user_inputs)
                            output_path = "temp_output.docx"
                            completed_doc.save(output_path)
                            st.session_state.generated_doc_path = output_path
                            st.session_state.generated_template_name = template_selection
                            st.success("✅ Document generated successfully!")
                        else:
                            st.warning("Please fill all required fields.")

            if "generated_doc_path" in st.session_state and os.path.exists(st.session_state.generated_doc_path):
                with open(st.session_state.generated_doc_path, "rb") as file:
                    st.download_button(
                        label="Download Completed Document",
                        data=file,
                        file_name=f"{st.session_state.generated_template_name}_completed.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
            else:
                st.write("You can directly download the template without editing.")
                with open(template_path, "rb") as file:
                    st.download_button(
                        label="Download Word Template",
                        data=file,
                        file_name=os.path.basename(template_path),
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

if __name__ == "__main__":
    main()
